import io
import datetime

from docx import Document  # https://python-docx.readthedocs.io/en/latest/
import srt  # https://srt.readthedocs.io/en/latest/



class Subs():
    def __init__(self):
        self.subs_dict = {}
        keys = ['index', 'start', 'end', 'fragment']
        for key in keys:
            self.subs_dict[key] = []
    
    
    def parse_time_code(self, time_code):
        seconds = 0
        for power, time_part in enumerate(reversed(time_code.split(':', 2) ) ):
            seconds += int(time_part) * 60 ** power
        return datetime.timedelta(seconds=seconds)
    
    
    def parse_time_interval(self, time_code_interval, seps):
        time_code_interval = time_code_interval.strip()
        if not time_code_interval:
            return None
        time_codes = []
        for sep in seps:
            if sep in time_code_interval:
                time_codes = time_code_interval.split(sep, 1)
                break
        if not time_codes:
            time_codes.append(time_code_interval)
        time_interval = []
        for time_code in time_codes:
            time_interval.append(self.parse_time_code(time_code.strip() ) )
        if len(time_interval) == 1:
            time_interval.append(None)
        return time_interval
    
    
    def time_remap(self, maping):
        remaped = []
        prev_start = None
        for start, end in maping:
            if prev_start is None:
                prev_start = start
                continue
            remaped.append([prev_start, start])
            prev_start = start
        
        for source_interval, dest_interval in zip(maping, remaped):
            source_end_specified = not source_interval[1] is None
            if source_end_specified:
                source_end_is_between_dests = \
                    dest_interval[0] < source_interval[1] < dest_interval[1]
            
            if source_end_specified and source_end_is_between_dests:
                dest_interval[1] = source_interval[1]
        return remaped
    
    
    def parse_ft_file(self, ft):
        try:
            doc = Document(ft)
        except Exception:
            raise Exception('f-err_file_format')
        
        try:
            material_number_string = doc.paragraphs[0].text.strip().split()[-1]
            self.material_number = int(material_number_string)
        except Exception:
            raise Exception('f-err_mat_num_format')
        
        duration_marker = 'Длительность'
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith(duration_marker):
                duration_time_code = paragraph.text.strip().split()[-1]
                break
        try:
            duration_time = self.parse_time_code(duration_time_code)
        except Exception:
            raise Exception('f-err_dur_format')
        
        try:
            table_of_fragments = doc.tables[1]
        except Exception:
            raise Exception('f-err_no_table')
        
        time_codes_col_index = 0
        # fragments in foreign language
        fragments_col_index = len(table_of_fragments.columns) - 1
        try:
            time_codes = table_of_fragments.column_cells(time_codes_col_index)
            fragments = table_of_fragments.column_cells(fragments_col_index)
        except Exception:
            raise Exception('f-err_no_col')
        
        time_title = time_codes.pop(0).text.strip()
        self.language_code = fragments.pop(0).text.strip().upper()
        
        sub_time_codes = []
        sub_fragments = []
        for time_code_cell, fragment_cell in zip(time_codes, fragments):
            time_code = time_code_cell.text.strip()
            fragment = fragment_cell.text.strip()
            if not time_code:
                continue
            if not fragment:
                raise Exception('f-err_empty_cell', 
                                self.language_code, 
                                time_code)
            sub_time_codes.append(time_code)
            sub_fragments.append(fragment)
        
        separators = ('-', '—', '–')
        time_mapping = []
        for time_code in sub_time_codes:
            try:
                time_mapping.append(
                    self.parse_time_interval(time_code, seps=separators))
            except Exception:
                raise Exception('f-err_time_format', time_title, time_code)
        time_mapping.append([duration_time, None])
        
        time_mapping = self.time_remap(time_mapping)
        
        for index, ((start, end), fragment) in enumerate(zip(time_mapping, 
                                                             sub_fragments)):
            self.subs_dict['index'].append(index + 1)
            self.subs_dict['start'].append(start)
            self.subs_dict['end'].append(end)
            self.subs_dict['fragment'].append(fragment)
    
    
    def compose_srt_file(self):
        sub_list = []
        for index in range(len(self.subs_dict['index'] ) ):
            sub_list.append(srt.Subtitle(self.subs_dict['index'][index],
                                         self.subs_dict['start'][index],
                                         self.subs_dict['end'][index],
                                         self.subs_dict['fragment'][index]))
        subs_bytes = srt.compose(sub_list, reindex=False).encode()
        self.subs_file = io.BytesIO(subs_bytes)
        self.subs_file.name = f'{self.material_number} '\
                               'Перевод Субтитров '\
                              f'{self.language_code}.srt'
        return self.subs_file
